import docx
from io import BytesIO
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
from PIL import Image, ImageDraw


document = docx.Document()

document.add_heading('Document Title', 0)
p = document.add_paragraph('How are u ren dong')
p.add_run(' bold').bold = True
p.add_run(' and some')
p.add_run(' italic.').italic = True
document.add_heading('level 1')

p = document.add_paragraph('How are u ren chu fei')

table = document.add_table(rows = 1, cols = 3, style = 'Table Grid')

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Gender'
hdr_cells[2].text = 'Age'
document.add_page_break()



p = document.add_paragraph()
r = p.add_run()
img_size = 20

for x in range(255):
    im = Image.new('RGB',(img_size, img_size), 'white')
    draw_obj = ImageDraw.Draw(im)
    draw_obj.ellipse((0,0,img_size-1,img_size-1),fill = 255 - x)
    fake_buf_file = BytesIO()
    im.save(fake_buf_file,"png")
    r.add_picture(fake_buf_file)
    fake_buf_file.close()

document.add_page_break()

p = document.add_paragraph()
text = '听闻原油期货要上市，扑克为大家准备了原油大礼包'

for i,ch in enumerate(text):
    run = p.add_run(ch)
    font = run.font
    font.name = 'Silom'
    font.size = Pt(16)
    font.color.rgb = RGBColor(i*10%200 + 55, i*20%200 + 55, i*30%200 + 55 )

document.save('test.docx')

